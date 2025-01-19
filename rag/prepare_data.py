from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.documents import Document
import pickle
from typing import List
from docx import Document as DocxDocument


def load_docx(file_path) -> List[Document]:
    """Load documents from file .docx and process paragraphs and tables."""
    doc = DocxDocument(file_path)
    content = []
    
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph element
            para = element.xpath(".//w:t")
            if para:
                text = ''.join([node.text for node in para if node.text])
                if text.strip():
                    content.append(text.strip())  # Each paragraph as a chunk
        elif element.tag.endswith('tbl'):  # Table element
            table_idx = [idx for idx, tbl in enumerate(doc.tables) if tbl._element == element]
            if table_idx:
                table = doc.tables[table_idx[0]]
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        content.append(' | '.join(row_data))  # Combine row cells into a single chunk
                        
    return [Document(page_content=c) for c in content]


def split_text_by_line(text: str) -> List[Document]:
    """Split text into chunks by line, each Document contains one line."""
    lines = text.split("\n")
    result = []
    for line in lines:
        line = line.strip()
        if line:  
            result.append(Document(page_content=line))
    return result


def save_2_pickle(documents: List[Document], file_path: str):
    """Save the list of Document objects to a pickle file."""
    with open(file_path, "wb") as file:
        pickle.dump(documents, file)


def load_data(file_path: str) -> List[Document]:
    with open(file_path, "rb") as file:
        documents = pickle.load(file)
    return documents


def prepare(file_path: str, file_destination_path: str):
    """Prepare and save data from .docx file into a pickle."""
    documents = load_docx(file_path)
    print(f"Loaded data from {file_path}")
    save_2_pickle(documents, file_destination_path)
    print(f"Saved to {file_destination_path}")


if __name__ == "__main__":
    
    prepare(
        "./data/Bộ luật-91-2015-QH13.docx",
        "./datapickle/dan_su.pkl"
    )
   
